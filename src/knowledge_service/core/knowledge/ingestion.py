"""ingestion.py

Purpose: Knowledge base population pipeline

Requirements:
--------------------------------------------------------------------------------
• Handle background processing
• Support multiple file formats
• Generate embeddings and store in ChromaDB

Key Components:
--------------------------------------------------------------------------------
  class KnowledgeIngester: ...
  @background_task def ingest_document(...)

Technology Stack:
--------------------------------------------------------------------------------
PyPDF2, python-docx, sentence-transformers

Core Design Principles:
--------------------------------------------------------------------------------
• Privacy-First: Sanitize all external-bound data
• Resilience: Implement retries and fallbacks
• Cost-Efficiency: Use semantic caching
• Extensibility: Use interfaces for pluggable components
• Observability: Add tracing spans for key operations
"""

import logging
import os
import uuid
from typing import Any, Dict, List, Optional

import chromadb
import pandas as pd
import pypdf
from chromadb.config import Settings
from docx import Document

from faultmaven.models import KnowledgeBaseDocument
from faultmaven.infrastructure.observability.tracing import trace
from faultmaven.infrastructure.security.redaction import DataSanitizer
from faultmaven.infrastructure.model_cache import model_cache


class KnowledgeIngester:
    """Handles asynchronous ingestion of documents into the knowledge base"""

    def __init__(self, chroma_persist_directory: str = "./chroma_db", settings=None):
        self.logger = logging.getLogger(__name__)
        self.sanitizer = DataSanitizer()
        
        # Get settings if not provided
        if settings is None:
            try:
                from faultmaven.config.settings import get_settings
                settings = get_settings()
            except Exception:
                settings = None
        
        # Initialize ChromaDB - default to K8s cluster for production-like development
        if settings:
            # Use settings-based configuration
            chromadb_url = settings.database.chromadb_url
            chromadb_host = settings.database.chromadb_host
            chromadb_port = settings.database.chromadb_port
            chromadb_auth_token = (
                settings.database.chromadb_auth_token.get_secret_value() 
                if settings.database.chromadb_auth_token 
                else "faultmaven-dev-chromadb-2025"
            )
        else:
            # No fallback - unified settings system is mandatory
            from faultmaven.models.exceptions import KnowledgeBaseError
            raise KnowledgeBaseError(
                "Knowledge ingestion requires unified settings system to be available",
                error_code="KNOWLEDGE_CONFIG_ERROR",
                context={"settings_available": settings is not None}
            )
        
        if chromadb_url:
            # Legacy URL-based configuration
            self.logger.info(f"Using ChromaDB HTTP client at {chromadb_url}")
            self.chroma_client = chromadb.HttpClient(
                host=chromadb_url.replace("http://", "")
                .replace("https://", "")
                .split(":")[0],
                port=int(chromadb_url.split(":")[-1]),
                settings=Settings(anonymized_telemetry=False, allow_reset=True),
            )
        elif chromadb_host != "localhost":
            # K8s cluster or external HTTP client (default)
            self.logger.info(f"Using ChromaDB HTTP client at {chromadb_host}:{chromadb_port}")
            self.chroma_client = chromadb.HttpClient(
                host=chromadb_host,
                port=chromadb_port,
                settings=Settings(
                    anonymized_telemetry=False,
                    allow_reset=True,
                    chroma_client_auth_provider="chromadb.auth.token_authn.TokenAuthClientProvider",
                    chroma_client_auth_credentials=chromadb_auth_token
                ),
            )
        else:
            # Local development with persistent client
            self.logger.info(
                f"Using ChromaDB PersistentClient at {chroma_persist_directory}"
            )
            self.chroma_client = chromadb.PersistentClient(
                path=chroma_persist_directory,
                settings=Settings(anonymized_telemetry=False, allow_reset=True),
            )

        # Get or create collection
        self.collection = self.chroma_client.get_or_create_collection(
            name="faultmaven_kb", metadata={"description": "FaultMaven Knowledge Base"}
        )

        # Initialize sentence transformer for embeddings using cached model
        self.embedding_model = model_cache.get_bge_m3_model()
        if self.embedding_model is None:
            self.logger.error("Failed to load BGE-M3 embedding model from cache")
            raise RuntimeError("BGE-M3 model unavailable - knowledge ingestion cannot proceed")
        else:
            self.logger.debug("Using cached BGE-M3 embedding model")

        # Supported file extensions
        self.supported_extensions = {
            ".txt": self._extract_text_txt,
            ".md": self._extract_text_txt,
            ".pdf": self._extract_text_pdf,
            ".docx": self._extract_text_docx,
            ".csv": self._extract_text_csv,
            ".json": self._extract_text_json,
            ".yaml": self._extract_text_yaml,
            ".yml": self._extract_text_yaml,
        }

    @trace("knowledge_base_ingest_document")
    async def ingest_document(
        self,
        file_path: str,
        title: str,
        document_type: str = "troubleshooting_guide",
        tags: Optional[List[str]] = None,
        source_url: Optional[str] = None,
        document_id: Optional[str] = None,
    ) -> str:
        """
        Ingest a document into the knowledge base (background task)

        Args:
            file_path: Path to the document file
            title: Document title
            document_type: Type of document
            tags: Optional tags for categorization
            source_url: Optional source URL
            document_id: Optional document ID to use (generates new if not provided)

        Returns:
            Document ID of the ingested document
        """
        if document_id is None:
            document_id = str(uuid.uuid4())

        try:
            self.logger.info(f"Starting ingestion of document: {title}")

            # Extract text content
            content = await self._extract_text(file_path)
            if not content:
                raise ValueError(f"Could not extract text from {file_path}")

            # Sanitize content
            sanitized_content = self.sanitizer.sanitize(content)

            # Create document object
            document = KnowledgeBaseDocument(
                document_id=document_id,
                title=title,
                content=sanitized_content,
                document_type=document_type,
                tags=tags or [],
                source_url=source_url,
            )

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {title}")
            return document_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {title}: {e}")
            raise

    async def _extract_text(self, file_path: str) -> str:
        """
        Extract text content from file based on its extension

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = os.path.splitext(file_path)[1].lower()

        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")

        extractor = self.supported_extensions[file_extension]
        return await extractor(file_path)

    async def _extract_text_txt(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    async def _extract_text_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from PDF {file_path}: {e}")
            raise

    async def _extract_text_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            self.logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise

    async def _extract_text_csv(self, file_path: str) -> str:
        """Extract text from CSV files"""
        try:
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            self.logger.error(f"Failed to extract text from CSV {file_path}: {e}")
            raise

    async def _extract_text_json(self, file_path: str) -> str:
        """Extract text from JSON files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import json

                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            self.logger.error(f"Failed to extract text from JSON {file_path}: {e}")
            raise

    async def _extract_text_yaml(self, file_path: str) -> str:
        """Extract text from YAML files"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                import yaml

                data = yaml.safe_load(f)
                return yaml.dump(data, default_flow_style=False)
        except Exception as e:
            self.logger.error(f"Failed to extract text from YAML {file_path}: {e}")
            raise

    async def _process_and_store(self, document: KnowledgeBaseDocument):
        """
        Process document content and store in ChromaDB

        Args:
            document: Document to process and store
        """
        # Split content into chunks
        chunks = self._split_content(document.content)

        # Generate embeddings for chunks
        embeddings = []
        for chunk in chunks:
            embedding = self.embedding_model.encode(chunk)
            embeddings.append(embedding.tolist())

        # Prepare metadata for each chunk
        metadatas = []
        ids = []

        for i, chunk in enumerate(chunks):
            chunk_id = f"{document.document_id}_chunk_{i}"
            metadata = {
                "document_id": document.document_id,
                "title": document.title,
                "document_type": document.document_type,
                "tags": ",".join(document.tags) if document.tags else "",
                "source_url": document.source_url or "",
                "chunk_index": i,
                "total_chunks": len(chunks),
                "created_at": document.created_at.isoformat(),
            }

            ids.append(chunk_id)
            metadatas.append(metadata)

        # Store in ChromaDB
        self.collection.add(
            embeddings=embeddings, documents=chunks, metadatas=metadatas, ids=ids
        )

        self.logger.info(
            f"Stored {len(chunks)} chunks for document {document.document_id}"
        )

    def _split_content(
        self, content: str, chunk_size: int = 1000, overlap: int = 200
    ) -> List[str]:
        """
        Split content into overlapping chunks

        Args:
            content: Content to split
            chunk_size: Maximum size of each chunk
            overlap: Overlap between chunks

        Returns:
            List of content chunks
        """
        if len(content) <= chunk_size:
            return [content]

        chunks = []
        start = 0

        while start < len(content):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(content):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if content[i] in ".!?":
                        end = i + 1
                        break

            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap
            if start >= len(content):
                break

        return chunks

    @trace("knowledge_base_search")
    async def search(
        self,
        query: str,
        n_results: int = 5,
        filter_metadata: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """
        Search the knowledge base

        Args:
            query: Search query
            n_results: Number of results to return
            filter_metadata: Optional metadata filters

        Returns:
            List of search results with documents and metadata
        """
        try:
            # Generate query embedding
            query_embedding = self.embedding_model.encode(query).tolist()

            # Prepare where clause for filtering
            where_clause = None
            if filter_metadata:
                where_clause = filter_metadata

            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                where=where_clause,
                include=["documents", "metadatas", "distances"],
            )

            # Format results
            formatted_results = []
            if results["documents"] and results["documents"][0]:
                for i, doc in enumerate(results["documents"][0]):
                    result = {
                        "document": doc,
                        "metadata": results["metadatas"][0][i],
                        "distance": results["distances"][0][i],
                        "relevance_score": 1
                        - results["distances"][0][i],  # Convert distance to relevance
                    }
                    formatted_results.append(result)

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []

    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document and all its chunks from the knowledge base

        Args:
            document_id: ID of the document to delete

        Returns:
            True if deletion was successful, False if document not found
        """
        try:
            # Find all chunks for this document
            self.logger.info(f"Attempting to delete document {document_id}")

            # First, let's see what's in the collection
            all_results = self.collection.get(include=["metadatas"], limit=10)
            self.logger.info(
                f"Sample collection contents: {len(all_results.get('metadatas', []))} items"
            )
            if all_results.get("metadatas"):
                for i, meta in enumerate(all_results["metadatas"][:3]):
                    self.logger.info(
                        f"Sample item {i}: {meta.get('document_id', 'no_id')}"
                    )

            results = self.collection.get(
                where={"document_id": document_id}, include=["metadatas"]
            )

            # ChromaDB returns IDs by default in results
            chunk_ids = results.get("ids", [])
            self.logger.info(
                f"Query results for {document_id}: found {len(chunk_ids)} chunk IDs"
            )

            if chunk_ids and len(chunk_ids) > 0:
                # Delete all chunks
                self.collection.delete(ids=chunk_ids)
                self.logger.info(
                    f"Deleted {len(chunk_ids)} chunks for document {document_id}"
                )
                return True
            else:
                self.logger.warning(f"No chunks found for document {document_id}")
                return False

        except Exception as e:
            self.logger.error(f"Failed to delete document {document_id}: {e}")
            return False

    def get_collection_stats(self) -> Dict[str, Any]:
        """
        Get statistics about the knowledge base collection

        Returns:
            Dictionary with collection statistics
        """
        try:
            count = self.collection.count()

            # Get sample of documents to analyze
            sample = self.collection.get(limit=1000, include=["metadatas"])

            # Analyze document types
            doc_types = {}
            tags = {}

            if sample["metadatas"]:
                for metadata in sample["metadatas"]:
                    doc_type = metadata.get("document_type", "unknown")
                    doc_types[doc_type] = doc_types.get(doc_type, 0) + 1

                    tag_list = metadata.get("tags", "").split(",")
                    for tag in tag_list:
                        tag = tag.strip()
                        if tag:
                            tags[tag] = tags.get(tag, 0) + 1

            return {
                "total_chunks": count,
                "document_types": doc_types,
                "top_tags": dict(
                    sorted(tags.items(), key=lambda x: x[1], reverse=True)[:10]
                ),
                "collection_name": self.collection.name,
            }

        except Exception as e:
            self.logger.error(f"Failed to get collection stats: {e}")
            return {}

    async def ingest_document_object(
        self,
        document: KnowledgeBaseDocument,
    ) -> str:
        """
        Ingest a document object into the knowledge base (for API uploads)

        Args:
            document: KnowledgeBaseDocument object with content already loaded

        Returns:
            Job ID for tracking the ingestion process
        """
        try:
            self.logger.info(f"Starting ingestion of document: {document.title}")

            # Sanitize content (already done in API, but double-check)
            sanitized_content = self.sanitizer.sanitize(document.content)

            # Update document with sanitized content
            document.content = sanitized_content

            # Process and store in chunks
            await self._process_and_store(document)

            self.logger.info(f"Successfully ingested document: {document.title}")

            # Generate job ID for tracking (in a real system, this would be stored)
            job_id = f"job_{document.document_id}"
            return job_id

        except Exception as e:
            self.logger.error(f"Failed to ingest document {document.title}: {e}")
            raise

    async def get_job_status(self, job_id: str) -> Optional[Dict[str, Any]]:
        """
        Get the status of an ingestion job

        Args:
            job_id: Job identifier

        Returns:
            Job status information or None if not found
        """
        # For now, return a simple completed status
        # In a real system, this would track actual job progress
        if job_id.startswith("job_"):
            document_id = job_id.replace("job_", "")
            return {
                "job_id": job_id,
                "document_id": document_id,
                "status": "completed",
                "progress": 100,
                "created_at": "2025-01-01T00:00:00",
                "completed_at": "2025-01-01T00:00:01",
            }
        return None

    async def list_documents(
        self,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 50,
        offset: int = 0,
    ) -> List[KnowledgeBaseDocument]:
        """
        List documents in the knowledge base

        Args:
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of documents
            offset: Number of documents to skip

        Returns:
            List of documents
        """
        try:
            # Build where clause for filtering
            where_clause = {}
            if document_type:
                where_clause["document_type"] = document_type
            if tags:
                # For simplicity, just filter by first tag
                # In a real implementation, would need more complex tag filtering
                where_clause["tags"] = {"$contains": tags[0]}

            # Get documents from ChromaDB
            results = self.collection.get(
                where=where_clause if where_clause else None,
                include=["metadatas"],
                limit=limit,
                offset=offset,
            )

            # Convert to document objects
            documents = []
            seen_doc_ids = set()

            if results["metadatas"]:
                for metadata in results["metadatas"]:
                    doc_id = metadata.get("document_id")
                    if doc_id and doc_id not in seen_doc_ids:
                        seen_doc_ids.add(doc_id)
                        # Create document object from metadata
                        doc = KnowledgeBaseDocument(
                            document_id=doc_id,
                            title=metadata.get("title", ""),
                            content="",  # Don't include full content in list
                            document_type=metadata.get("document_type", ""),
                            tags=(
                                metadata.get("tags", "").split(",")
                                if metadata.get("tags")
                                else []
                            ),
                            source_url=metadata.get("source_url"),
                        )
                        documents.append(doc)

            return documents

        except Exception as e:
            self.logger.error(f"Failed to list documents: {e}")
            return []

    async def get_document(self, document_id: str) -> Optional[KnowledgeBaseDocument]:
        """
        Get a specific document by ID

        Args:
            document_id: Document identifier

        Returns:
            Document object or None if not found
        """
        try:
            # Get all chunks for this document
            results = self.collection.get(
                where={"document_id": document_id}, include=["documents", "metadatas"]
            )

            if not results["documents"] or not results["documents"]:
                return None

            # Reconstruct document from chunks
            chunks = results["documents"]
            metadata = results["metadatas"][0] if results["metadatas"] else {}

            # Combine all chunks to reconstruct content
            content = " ".join(chunks)

            doc = KnowledgeBaseDocument(
                document_id=document_id,
                title=metadata.get("title", ""),
                content=content,
                document_type=metadata.get("document_type", ""),
                tags=(
                    metadata.get("tags", "").split(",") if metadata.get("tags") else []
                ),
                source_url=metadata.get("source_url"),
            )

            return doc

        except Exception as e:
            self.logger.error(f"Failed to get document {document_id}: {e}")
            return None

    async def search_documents(
        self,
        query: str,
        document_type: Optional[str] = None,
        tags: Optional[List[str]] = None,
        limit: int = 10,
    ) -> List[Dict[str, Any]]:
        """
        Search documents and return results with scores

        Args:
            query: Search query
            document_type: Filter by document type
            tags: Filter by tags
            limit: Maximum number of results

        Returns:
            List of search results with document info and scores
        """
        try:
            # Build filter metadata
            filter_metadata = {}
            if document_type:
                filter_metadata["document_type"] = document_type
            if tags:
                filter_metadata["tags"] = {"$contains": tags[0]}

            # Search using existing search method
            results = await self.search(
                query=query,
                n_results=limit,
                filter_metadata=filter_metadata if filter_metadata else None,
            )

            # Format for API response
            formatted_results = []
            for result in results:
                metadata = result["metadata"]
                formatted_result = {
                    "document_id": metadata.get("document_id"),
                    "title": metadata.get("title"),
                    "document_type": metadata.get("document_type"),
                    "tags": (
                        metadata.get("tags", "").split(",")
                        if metadata.get("tags")
                        else []
                    ),
                    "score": result["relevance_score"],
                    "snippet": (
                        result["document"][:200] + "..."
                        if len(result["document"]) > 200
                        else result["document"]
                    ),
                }
                formatted_results.append(formatted_result)

            return formatted_results

        except Exception as e:
            self.logger.error(f"Search failed: {e}")
            return []
